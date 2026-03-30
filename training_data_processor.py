"""
Training Data Processor for Phantom Visible Scripter
Extracts and processes user's existing scripts for style training
"""

import os
import json
from pathlib import Path
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class ScriptStyleProcessor:
    """Process existing scripts to extract writing style patterns"""
    
    def __init__(self, previous_work_dir: str = "previous_work"):
        self.previous_work_dir = Path(previous_work_dir)
        self.training_data = []
        
    def extract_scripts_from_docx(self) -> List[Dict[str, Any]]:
        """Extract scripts from .docx files in previous_work folder"""
        scripts = []
        
        if not self.previous_work_dir.exists():
            logger.warning(f"Previous work directory not found: {self.previous_work_dir}")
            return scripts
        
        # Find all .docx files
        docx_files = list(self.previous_work_dir.glob("*.docx"))
        
        for file_path in docx_files:
            try:
                script_data = self._process_docx_file(file_path)
                if script_data:
                    scripts.append(script_data)
                    logger.info(f"Processed: {file_path.name}")
            except Exception as e:
                logger.error(f"Error processing {file_path.name}: {e}")
        
        logger.info(f"Extracted {len(scripts)} scripts from {len(docx_files)} files")
        return scripts
    
    def _process_docx_file(self, file_path: Path) -> Dict[str, Any]:
        """Process individual .docx file to extract script content"""
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx not installed. Install with: pip install python-docx")
                return None
            
            # Extract text from document
            doc = Document(str(file_path))
            full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract metadata
            return {
                "filename": file_path.name,
                "content": full_text,
                "word_count": len(full_text.split()),
                "char_count": len(full_text),
                "paragraph_count": len(doc.paragraphs),
                "style_features": self._extract_style_features(full_text)
            }
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            return None
    
    def _extract_style_features(self, text: str) -> Dict[str, Any]:
        """Extract stylistic features from script text"""
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        return {
            "avg_sentence_length": sum(len(s.split()) for s in sentences) / len(sentences) if sentences else 0,
            "question_ratio": len([s for s in sentences if '?' in s]) / len(sentences) if sentences else 0,
            "exclamation_ratio": len([s for s in sentences if '!' in s]) / len(sentences) if sentences else 0,
            "contraction_count": len([word for s in sentences for word in s.split() if "'" in word]),
            "hook_patterns": self._find_hook_patterns(text),
            "transition_words": self._find_transition_words(text),
            "vocabulary_diversity": len(set(text.lower().split())) / len(text.split()) if text.split() else 0
        }
    
    def _find_hook_patterns(self, text: str) -> List[str]:
        """Find common hook patterns in script opening"""
        hooks = []
        lines = text.split('\n')[:5]  # Check first 5 lines
        
        for line in lines:
            line = line.strip()
            if any(pattern in line.lower() for pattern in [
                'what if', 'imagine', 'did you know', 'the truth is', 
                'have you ever', 'there\'s one thing', 'let me ask'
            ]):
                hooks.append(line)
        
        return hooks
    
    def _find_transition_words(self, text: str) -> List[str]:
        """Find transition words and phrases"""
        transition_patterns = [
            'now', 'but', 'so', 'however', 'meanwhile', 'furthermore',
            'on the other hand', 'let\'s talk about', 'moving on',
            'here\'s the thing', 'what\'s interesting', 'the key is'
        ]
        
        found = []
        text_lower = text.lower()
        for pattern in transition_patterns:
            if pattern in text_lower:
                found.append(pattern)
        
        return found
    
    def create_training_dataset(self, output_file: str = "training_data/style_dataset.json"):
        """Create training dataset from processed scripts"""
        scripts = self.extract_scripts_from_docx()
        
        if not scripts:
            logger.error("No scripts to process")
            return
        
        training_data = {
            "metadata": {
                "total_scripts": len(scripts),
                "total_words": sum(s["word_count"] for s in scripts),
                "avg_word_count": sum(s["word_count"] for s in scripts) / len(scripts),
                "style_patterns": self._analyze_global_patterns(scripts)
            },
            "scripts": scripts
        }
        
        # Save training dataset
        output_path = Path(output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(training_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Training dataset saved to: {output_path}")
        return training_data
    
    def _analyze_global_patterns(self, scripts: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze patterns across all scripts"""
        all_hooks = []
        all_transitions = []
        all_vocab_sizes = []
        
        for script in scripts:
            features = script.get("style_features", {})
            all_hooks.extend(features.get("hook_patterns", []))
            all_transitions.extend(features.get("transition_words", []))
            all_vocab_sizes.append(features.get("vocabulary_diversity", 0))
        
        return {
            "common_hooks": self._most_common(all_hooks),
            "common_transitions": self._most_common(all_transitions),
            "avg_vocabulary_diversity": sum(all_vocab_sizes) / len(all_vocab_sizes) if all_vocab_sizes else 0
        }
    
    def _most_common(self, items: List[str]) -> List[str]:
        """Find most common items in list"""
        from collections import Counter
        counter = Counter(items)
        return [item for item, count in counter.most_common(5)]

def main():
    """Main function to process previous work"""
    processor = ScriptStyleProcessor()
    
    print("🎯 Processing your previous scripts for style training...")
    dataset = processor.create_training_dataset()
    
    if dataset:
        print(f"✅ Processed {dataset['metadata']['total_scripts']} scripts")
        print(f"📊 Average word count: {dataset['metadata']['avg_word_count']:.0f}")
        print(f"🎭 Common hooks: {', '.join(dataset['metadata']['style_patterns']['common_hooks'][:3])}")
        print(f"🔄 Common transitions: {', '.join(dataset['metadata']['style_patterns']['common_transitions'][:3])}")
        print("\n📁 Training data saved to: training_data/style_dataset.json")
        print("\n🚀 Next steps:")
        print("1. Install python-docx: pip install python-docx")
        print("2. Use this dataset to fine-tune your model")
        print("3. Update ScriptingAgent to use your trained style")

if __name__ == "__main__":
    main()
